from pathlib import Path
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "documents" / "GENEVIEVE_HEALTH_Letter_to_Irene_Program_and_Trial_Guide_V6.docx"
GA = ROOT / "public" / "demo" / "assets" / "icon-512.png"
TREE = ROOT / "public" / "demo" / "assets" / "genevieve-tree-logo-approved-original.jpeg"

INK = "171217"
GOLD = "D9B83F"
GOLD_DARK = "8B6816"
PINK = "F8E8EF"
ROSE = "8C4567"
PAPER = "FFFAFB"
GREY = "705F68"
RED = "BF1E2E"
GREEN = "21865B"
WHITE = "FFFFFF"


def set_cell_fill(cell, color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), color)


def set_cell_margins(cell, top=120, start=150, bottom=120, end=150):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def keep_with_next(paragraph):
    paragraph.paragraph_format.keep_with_next = True


def page_field(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("GENEVIEVE HEALTH™  •  IRENE TRIAL GUIDE  •  ")
    run.font.name = "Calibri"
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor.from_string(GREY)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    paragraph._p.append(fld)


def style_doc(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.60)
    section.bottom_margin = Inches(0.58)
    section.left_margin = Inches(0.82)
    section.right_margin = Inches(0.82)
    section.header_distance = Inches(0.32)
    section.footer_distance = Inches(0.30)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(9.8)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(4.5)
    normal.paragraph_format.line_spacing = 1.12

    for name, size, before, after, color in (
        ("Title", 25, 0, 12, INK),
        ("Heading 1", 16.5, 12, 7, INK),
        ("Heading 2", 12.5, 9, 5, ROSE),
        ("Heading 3", 10.5, 7, 3, GOLD_DARK),
    ):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for section in doc.sections:
        page_field(section.footer.paragraphs[0])


def brand_header(doc):
    table = doc.add_table(rows=1, cols=3)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    widths = [Inches(0.82), Inches(4.95), Inches(0.82)]
    for cell, width in zip(table.rows[0].cells, widths):
        cell.width = width
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_margins(cell, 35, 80, 35, 80)
    set_cell_fill(table.cell(0, 1), INK)
    left = table.cell(0, 0).paragraphs[0]
    left.alignment = WD_ALIGN_PARAGRAPH.CENTER
    left.add_run().add_picture(str(GA), width=Inches(0.54))
    centre = table.cell(0, 1).paragraphs[0]
    centre.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = centre.add_run("GENEVIEVE App™\n")
    run.bold = True
    run.font.size = Pt(13.5)
    run.font.color.rgb = RGBColor.from_string(WHITE)
    run = centre.add_run("GENEVIEVE HEALTH™  |  MENTAL HEALTH PRACTICE SUPPORT")
    run.bold = True
    run.font.size = Pt(7.3)
    run.font.color.rgb = RGBColor.from_string(GOLD)
    right = table.cell(0, 2).paragraphs[0]
    right.alignment = WD_ALIGN_PARAGRAPH.CENTER
    right.add_run().add_picture(str(TREE), width=Inches(0.50))
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("SAFETY FROM ROOTS TO EVERY JOURNEY.")
    r.bold = True
    r.font.size = Pt(7)
    r.font.color.rgb = RGBColor.from_string(GOLD_DARK)
    r.font.letter_spacing = Pt(0.7)


def callout(doc, label, text, fill=PINK, border=GOLD_DARK):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_fill(cell, fill)
    set_cell_margins(cell, 105, 160, 105, 160)
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        tag = OxmlElement(f"w:{edge}")
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), "9")
        tag.set(qn("w:color"), border)
        borders.append(tag)
    tc_pr.append(borders)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(label.upper() + "  ")
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(border)
    r = p.add_run(text)
    r.font.color.rgb = RGBColor.from_string(INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def heading(doc, text, level=1):
    return doc.add_heading(text, level=level)


def bullet(doc, text, level=0, bold_lead=None):
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    p.paragraph_format.left_indent = Inches(0.28 + 0.22 * level)
    p.paragraph_format.first_line_indent = Inches(-0.18)
    p.paragraph_format.space_after = Pt(3)
    if bold_lead and text.startswith(bold_lead):
        r = p.add_run(bold_lead)
        r.bold = True
        p.add_run(text[len(bold_lead):])
    else:
        p.add_run(text)
    return p


def numbered_step(doc, number, title, body, owner=None):
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(0.55)
    table.columns[1].width = Inches(5.95)
    number_cell, text_cell = table.rows[0].cells
    set_cell_fill(number_cell, INK)
    set_cell_fill(text_cell, PAPER)
    set_cell_margins(number_cell, 75, 70, 75, 70)
    set_cell_margins(text_cell, 70, 140, 70, 140)
    number_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = number_cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(str(number))
    r.bold = True
    r.font.size = Pt(13)
    r.font.color.rgb = RGBColor.from_string(GOLD)
    p = text_cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor.from_string(ROSE)
    p = text_cell.add_paragraph(body)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.04
    for run in p.runs:
        run.font.size = Pt(9.1)
    if owner:
        p = text_cell.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run("Owner: ")
        r.bold = True
        r.font.color.rgb = RGBColor.from_string(GOLD_DARK)
        own = p.add_run(owner)
        own.font.size = Pt(8.6)
        for run in p.runs:
            run.font.size = Pt(8.6)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def page_break(doc):
    doc.add_page_break()


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for i, (header, width) in enumerate(zip(headers, widths)):
        table.columns[i].width = Inches(width)
        cell = table.rows[0].cells[i]
        set_cell_fill(cell, INK)
        set_cell_margins(cell, 80, 100, 80, 100)
        p = cell.paragraphs[0]
        r = p.add_run(header)
        r.bold = True
        r.font.size = Pt(7.6)
        r.font.color.rgb = RGBColor.from_string(GOLD)
    set_repeat_table_header(table.rows[0])
    for index, row in enumerate(rows):
        cells = table.add_row().cells
        for i, (value, width) in enumerate(zip(row, widths)):
            cells[i].width = Inches(width)
            cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cells[i], 65, 100, 65, 100)
            if index % 2 == 0:
                set_cell_fill(cells[i], "FFF6F9")
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.08
            r = p.add_run(value)
            r.font.size = Pt(7.7)
    return table


def build():
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    style_doc(doc)
    brand_header(doc)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run("18 July 2026")
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(GOLD_DARK)
    doc.add_paragraph("Private and confidential\nIrene\nPractice Director and Clinical Psychologist\nMood & Mind Centre")
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("A LETTER TO IRENE")
    r.bold = True
    r.font.size = Pt(8)
    r.font.color.rgb = RGBColor.from_string(ROSE)
    r.font.letter_spacing = Pt(1.2)
    title = doc.add_paragraph(style="Title")
    title.add_run("Your GENEVIEVE HEALTH™ Psychology Practice Safety Program and Trial Guide")
    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(14)
    r = subtitle.add_run("A connected, human-led operating system for Irene, clinicians, reception and approved governance roles")
    r.italic = True
    r.font.size = Pt(11.5)
    r.font.color.rgb = RGBColor.from_string(GREY)

    doc.add_paragraph("Dear Irene,")
    doc.add_paragraph("I am pleased to place this trial version of the GENEVIEVE HEALTH™ Psychology Practice Safety Program in your hands. It brings your director dashboard, each worker’s private phone view and a dedicated reception base into one connected operational system—while preserving clear role boundaries and keeping clinical judgement with qualified people.")
    doc.add_paragraph("The program is designed to help the practice notice operational risk earlier, give staff a simple way to ask for support, and leave a clear record that an issue was owned and acted upon. It does not diagnose, score a client’s clinical risk, replace professional judgement or contact emergency services or third parties automatically.")
    callout(doc, "Trial commitment", "A safety alert cannot be deactivated merely because it was seen or a message was sent. The completed action must be recorded, then Irene, an authorised supervisor, clinical lead or specifically permitted safety person must sign it off before closure.", "FFF6DF", GOLD_DARK)
    doc.add_paragraph("The pages that follow explain what the program does, who can see what, how the workload protections operate, and the exact steps for running the controlled trial this week.")
    doc.add_paragraph("Yours sincerely,")
    p = doc.add_paragraph()
    r = p.add_run("Tracey Ann Kennedy")
    r.bold = True
    r.font.size = Pt(11.5)
    r.font.color.rgb = RGBColor.from_string(ROSE)
    doc.add_paragraph("Founder, GENEVIEVE App™\nGENEVIEVE HEALTH™")
    callout(doc, "Controlled trial", "Use fictional or coded information only until Irene has approved the users, roles, procedures and production settings.", "FDECEF", RED)

    page_break(doc)
    brand_header(doc)
    heading(doc, "1. What the program does")
    doc.add_paragraph("GENEVIEVE HEALTH™ connects the practice’s operational safety work without giving every person access to everything. Each base is designed around the minimum information its user needs.")
    add_table(doc, ["BASE", "WHAT IT DOES", "PRIVACY BOUNDARY"], [
        ("Irene’s Connected Dashboard", "Whole-practice safety alerts, workload profiles, staff access, assignments, communication, action records and supervisor closure.", "Irene and approved oversight roles only."),
        ("Staff Phone App", "Own messages, tasks, shifts, lunch protection, workload check-ins, coded schedule mix and age-pattern preferences.", "Only that worker’s operational information."),
        ("Reception Base", "Coded callbacks, missed contacts, approved scripts, transfers, Irene contact, lunch and support controls.", "No therapy notes, diagnosis or clinical judgement."),
        ("Mother Board / Governance", "Only the oversight, audit and whole-practice information Irene has expressly authorised.", "Blocked until Irene approves the exact account."),
        ("Practice Safety Dashboard", "Existing continuity, supervision, emergency cover, WHS, document and evidence controls.", "Permissions still apply to each module."),
    ], [1.35, 3.15, 2.0])
    heading(doc, "Designed for phones and desktop", 2)
    bullet(doc, "Staff and reception can add the program to an iPhone or Android home screen as a progressive web app.")
    bullet(doc, "The last open page is remembered. Browser Back and Forward return to the previous program page.")
    bullet(doc, "Pages, forms, links, open/close controls and the GENEVIEVE Safety Guide are keyboard- and touch-operable.")
    bullet(doc, "Two-way messages and actions return to Irene’s dashboard; staff do not see each other’s private requests.")
    heading(doc, "The GENEVIEVE Safety Guide", 2)
    doc.add_paragraph("The in-program guide answers common operational questions and opens the correct page. It is deliberately rule-based: it does not diagnose, produce clinical risk scores or independently call anyone. It routes the user to an authorised human and the practice’s approved procedure.")
    callout(doc, "Emergency boundary", "When immediate emergency assistance is needed, follow the practice’s approved emergency procedure and call 000. The program does not contact emergency services automatically.", "FFF0F1", RED)

    page_break(doc)
    brand_header(doc)
    heading(doc, "2. The safety-alert control")
    doc.add_paragraph("The alert workflow is designed to prevent a warning disappearing before the protective action has definitely occurred.")
    numbered_step(doc, 1, "Detect and display", "The system checks active shifts and coded schedules for lunch timing, approved hours, staff support requests, session limits, high-support load and short adult/child transition buffers.")
    numbered_step(doc, 2, "Escalate", "Amber indicates attention is approaching or needed. A lunch more than 15 minutes overdue, an hours-limit breach, a high workload/support request, or an approved workload-limit breach escalates to a red human review.")
    numbered_step(doc, 3, "Acknowledge", "The worker or authorised supervisor records that the alert has been seen. Acknowledgement alone does not close or deactivate it.")
    numbered_step(doc, 4, "Record the action actually completed", "Record the real protective action—for example: coverage arranged, lunch commenced, schedule rebalanced, recovery buffer added, debrief booked, work transferred or hours ended.")
    numbered_step(doc, 5, "Supervisor sign-off and close", "Only Irene, an authorised supervisor, clinical lead or person holding explicit safety-sign-off permission can close the alert, and only after an action record exists.")
    add_table(doc, ["ALERT", "DEFAULT RESPONSE", "CLOSURE EVIDENCE"], [
        ("Lunch due / overdue", "Confirm coverage and break timing; escalate after the approved threshold.", "Break action recorded + authorised sign-off."),
        ("Hours limit", "Stop, transfer or reduce work; review the cause with the worker.", "Hours/work action + supervisor sign-off."),
        ("High workload / support request", "Prompt private contact, relief, debrief or workload adjustment.", "Human contact/action + sign-off."),
        ("Session or high-support load", "Rebalance, add recovery time or provide additional support.", "Schedule response + sign-off."),
        ("Adult/child transition", "Add the worker’s approved transition buffer or reorganise the day.", "Buffer/schedule action + sign-off."),
    ], [1.5, 3.15, 1.85])

    page_break(doc)
    brand_header(doc)
    heading(doc, "3. Workload protection for psychologists and reception")
    doc.add_paragraph("Psychology work can involve sustained cognitive and emotional demand. Moving rapidly between adult and child work may also require a significant change in pace, communication style, materials and emotional regulation. GENEVIEVE therefore treats the mix and sequencing of work as a consultation and planning matter—not a clinical rating of clients or a judgement about the worker.")
    heading(doc, "Controls Irene can agree with each worker", 2)
    bullet(doc, "Maximum daily hours and maximum daily sessions.", bold_lead="Maximum daily hours")
    bullet(doc, "Maximum number of high-support sessions in a day.", bold_lead="Maximum number")
    bullet(doc, "Lunch-due timing and protected coverage.", bold_lead="Lunch-due timing")
    bullet(doc, "Minimum recovery buffer when the schedule changes between adult and child/adolescent work.", bold_lead="Minimum recovery buffer")
    bullet(doc, "Preferred age pattern: mixed work with buffers, children-focused days, adult-focused days, or a custom agreed pattern.", bold_lead="Preferred age pattern")
    bullet(doc, "Private workload check-in, coverage request, debrief preference and follow-up action.", bold_lead="Private workload check-in")
    heading(doc, "How the system helps", 2)
    add_table(doc, ["BEFORE THE DAY", "DURING THE DAY", "AFTER AN ALERT"], [
        ("Load only coded schedule information. Review session count, age mix, high-support count and required buffers.", "Worker starts the shift. The lunch timer and hours protection run. The worker can privately report stretching/high demand.", "Irene or the authorised owner records the intervention, confirms recovery/coverage and signs off only when satisfied."),
    ], [2.15, 2.15, 2.2])
    callout(doc, "Important", "Do not use the high-support or age-mix fields as a clinical risk score, productivity measure, disciplinary score or substitute for consultation. Limits must be set and reviewed with workers and adapted to the actual practice.", "FFF6DF", GOLD_DARK)
    heading(doc, "Reception must be protected too", 2)
    doc.add_paragraph("Reception workers may face competing demands, distressed contacts, privacy pressure and interrupted breaks. Their base includes the same shift, lunch, workload and private support controls. They use approved words, keep only minimum coded operational information, and transfer clinical judgement to an authorised clinician.")

    page_break(doc)
    brand_header(doc)
    heading(doc, "4. Step-by-step: set up the controlled trial")
    numbered_step(doc, 1, "Irene signs in first", "Open Irene’s Connected Dashboard and use Sign in with ChatGPT. The first authorised live account becomes the director account. Confirm the displayed name and role.", "Irene")
    numbered_step(doc, 2, "Authorise exact staff accounts", "Open Staff Access. Enter the exact email each person will use for ChatGPT sign-in and choose the correct role. Give Mother Board access only to the specific people Irene approves.", "Irene")
    numbered_step(doc, 3, "Confirm role boundaries", "Open Privacy Mother Board and review what each role receives. Reception must not receive clinical notes; staff must not see another worker’s private support request.", "Irene + privacy lead")
    numbered_step(doc, 4, "Agree workload profiles", "Open Workload Protection. Consult each worker and set daily hours, sessions, high-support limit, lunch timing and an age-transition buffer. Review the worker’s preferred age pattern and approve it only after agreement.", "Irene / authorised supervisor + worker")
    numbered_step(doc, 5, "Prepare a coded trial schedule", "Use fictional codes only. Include a small mix of adult, child/adolescent and high-support blocks so the team can see the workload and transition controls operate without using real client information.", "Irene / authorised scheduler")
    numbered_step(doc, 6, "Prepare reception", "Authorise the reception account, open Reception Base, read each approved script, confirm the escalation owner and practise accepting, completing and escalating fictional coded items.", "Reception lead + Irene")
    numbered_step(doc, 7, "Confirm emergency and escalation procedures", "Place the practice’s approved emergency procedure beside the trial. Confirm who owns red alerts, backup ownership, expected response time and who has safety-sign-off permission.", "Irene")
    callout(doc, "Access setting", "Keep the trial owner-only until Irene has approved the exact staff identities and roles. Do not broaden access simply to make testing easier.", "FDECEF", RED)

    page_break(doc)
    brand_header(doc)
    heading(doc, "5. Step-by-step: run a normal trial day")
    numbered_step(doc, 1, "Start the shift", "Each worker opens My Safety and selects Start shift. This activates hours and lunch protection for that worker.", "Each worker")
    numbered_step(doc, 2, "Review today’s work mix", "Psychologists view their coded age and support mix. If the pattern feels tiring or does not match the agreed preference, they request a buffer, coverage or schedule review before the day compounds.", "Worker + scheduler")
    numbered_step(doc, 3, "Use private two-way communication", "Staff or reception send Irene a short operational message. Do not include client-identifying, therapy or health information. Irene replies through Staff Communications.", "All authorised users")
    numbered_step(doc, 4, "Protect lunch", "When lunch is due, arrange coverage and select Start lunch. The reminder becomes urgent if the approved overdue threshold is exceeded. Select Finish lunch only after the break is completed.", "Worker + coverage owner")
    numbered_step(doc, 5, "Respond to a safety alert", "Acknowledge it, speak with the worker, complete the protective action and record what actually happened. Do not close it to clear the screen.", "Alert owner")
    numbered_step(doc, 6, "Complete supervisor sign-off", "When the action record is accurate and the issue has been addressed, Irene or another authorised sign-off person selects Supervisor sign-off and close.", "Irene / authorised sign-off person")
    numbered_step(doc, 7, "Close the day", "End active shifts, finish coded reception items, review open red and amber alerts, hand over unresolved actions and note improvements for the trial review.", "Irene + closing owner")
    callout(doc, "Back pages and memory", "The program remembers the last page used on each base. Browser Back and Forward restore prior pages; normal navigation adds a page to that history. The Safety Guide can be opened and closed at any time.", "FFF6DF", GOLD_DARK)

    page_break(doc)
    brand_header(doc)
    heading(doc, "6. Reception quick guide")
    add_table(doc, ["SITUATION", "RECEPTION ACTION", "ESCALATE WHEN"], [
        ("Callback requested", "Create a coded callback, confirm the approved contact channel, accept ownership and record the operational outcome.", "The approved deadline or attempt threshold is reached."),
        ("Clinician unavailable", "Use the approved unavailability script and offer an authorised callback pathway.", "An authorised owner is not available within the approved timeframe."),
        ("Distressed or concerning contact", "Use the approved transfer script. Connect to an authorised clinician or the practice’s approved emergency pathway. Do not assess clinical risk.", "Promptly—according to Irene’s procedure and human escalation chain."),
        ("Missed contact", "Keep the item open, record the attempt and the next owner. Use minimum coded information.", "The practice’s approved threshold is reached or ownership is unclear."),
        ("Reception workload high", "Open My Safety, request coverage, a break, debrief or private Irene contact.", "The worker reports high demand, support is requested or a protected break is missed."),
    ], [1.45, 3.15, 1.9])
    heading(doc, "Approved words built into the base", 2)
    for title, text in (
        ("Clinician unavailable", "The clinician is unavailable. I can record the minimum information needed and arrange an authorised callback."),
        ("Concerning contact", "I am going to connect you with an authorised clinician or the practice’s approved emergency pathway. Reception does not make a clinical assessment."),
        ("Voicemail privacy", "Before leaving a message, confirm the approved contact channel and whether details may be left."),
        ("Immediate emergency", "If immediate emergency assistance is required, call 000. This program does not contact emergency services automatically."),
    ):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.18)
        p.paragraph_format.first_line_indent = Inches(-0.18)
        r = p.add_run(title + " — ")
        r.bold = True
        r.font.color.rgb = RGBColor.from_string(ROSE)
        p.add_run(text)
    callout(doc, "Reception boundary", "Record only the minimum operational information. Do not copy therapy notes, diagnoses, health histories or unnecessary identifying information into the reception queue or messages.", "FDECEF", RED)

    page_break(doc)
    brand_header(doc)
    heading(doc, "7. Trial acceptance checklist")
    checks = [
        "Irene can sign in, open every director page and use Back/Forward page history.",
        "Only exact authorised accounts receive a role; unapproved accounts are blocked.",
        "Staff and reception see only their own or expressly assigned operational information.",
        "Messages move both ways between each private base and Irene’s dashboard.",
        "Reception can accept, complete and escalate a coded item; escalation appears for Irene.",
        "Lunch timing produces an amber reminder and escalates red after the approved overdue threshold.",
        "An alert cannot close without a completed action record and authorised supervisor sign-off.",
        "Hours, session count, high-support load and short adult/child transition buffers display correctly.",
        "A worker can send an age-pattern preference for supervisor review and approval.",
        "The Safety Guide opens, closes and routes each user to the correct operational page.",
        "The phone app installs or can be added to the home screen on the trial iPhone and Android device.",
        "Only fictional/coded data was used and no emergency-service call was expected from the program.",
    ]
    for item in checks:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.34)
        p.paragraph_format.first_line_indent = Inches(-0.34)
        p.paragraph_format.space_after = Pt(5)
        r = p.add_run("☐  ")
        r.bold = True
        r.font.size = Pt(12)
        r.font.color.rgb = RGBColor.from_string(GOLD_DARK)
        p.add_run(item)
    heading(doc, "Before real-person production use", 2)
    doc.add_paragraph("Complete privacy and security review, staff consultation, emergency-procedure approval, retention settings, production monitoring, incident response, backup arrangements, device testing and any legal/professional review Irene requires. The controlled trial is a working demonstration, not evidence that every production obligation has been completed.")
    heading(doc, "Reference framework", 2)
    sources = [
        "Safe Work Australia — Psychosocial hazards; Job demands; Managing fatigue risks; Health care and social assistance: fatigue.",
        "Workplace Health and Safety Queensland — Managing the risk of psychosocial hazards at work Code of Practice 2022; High and/or low job demands.",
        "Australian Health Practitioner Regulation Agency — Shared Code of Conduct.",
    ]
    for source in sources:
        bullet(doc, source)
    callout(doc, "Human-led safety", "GENEVIEVE supports visibility, ownership and evidence. Irene and appropriately authorised people remain responsible for judgement, action, escalation and review.", "FFF6DF", GOLD_DARK)

    props = doc.core_properties
    props.title = "GENEVIEVE HEALTH — Letter to Irene: Psychology Practice Safety Program and Trial Guide V6"
    props.subject = "Professional program overview and step-by-step controlled trial guide"
    props.author = "GENEVIEVE App™"
    props.keywords = "GENEVIEVE HEALTH, Irene, psychology practice safety, reception, workload, trial guide"
    for index, shape in enumerate(doc.inline_shapes):
        description = "Approved GENEVIEVE App GA emblem" if index % 2 == 0 else "Approved GENEVIEVE tree-and-roots mark"
        shape._inline.docPr.set("descr", description)
        shape._inline.docPr.set("title", description)
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
